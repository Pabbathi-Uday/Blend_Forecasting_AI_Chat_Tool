# Streamlit and UI
import streamlit as st
from streamlit_float import *
import pandas as pd
import config
import markdown2
from io import BytesIO
import markdown2
from xhtml2pdf import pisa
import os
from docx import Document as docx_document
import numpy as np
import pandasai as pai

from pandas.api.types import (
    is_categorical_dtype,
    is_datetime64_any_dtype,
    is_numeric_dtype,
    is_object_dtype,
)

# System and utilities
import os
import json
import docx
import re
import shutil
import traceback
from io import BytesIO
from datetime import datetime
from typing import List, Dict, Any, Union
from concurrent.futures import ThreadPoolExecutor, as_completed

# NLP and PDF tools
import nltk
nltk.download('punkt')
from fpdf import FPDF
from xhtml2pdf import pisa

# OpenAI and LangChain
import openai
from openai import OpenAI
from langchain import hub
from langchain.chat_models import ChatOpenAI
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate, ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser, JsonOutputParser
from langchain_core.output_parsers.string import StrOutputParser
from langchain.chains.query_constructor.schema import AttributeInfo
from langchain.chains.query_constructor.base import (
    StructuredQueryOutputParser,
    get_query_constructor_prompt,
)
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
#from langchain.vectorstores import Chroma
from langchain_community.vectorstores.chroma import Chroma 
from langchain.retrievers import EnsembleRetriever
from langchain.retrievers.self_query.base import SelfQueryRetriever
from langchain.retrievers.self_query.chroma import ChromaTranslator
from langchain.schema.runnable import RunnablePassthrough, RunnableParallel, RunnableLambda

# Evaluation
from deepeval.metrics import (
    AnswerRelevancyMetric,
    FaithfulnessMetric,
    ContextualPrecisionMetric,
    ContextualRecallMetric
)
from deepeval.test_case import LLMTestCase

# Statistical packages
import statsmodels.api as sm
from statsmodels.stats.diagnostic import het_breuschpagan
from statsmodels.stats.stattools import durbin_watson

# Prompt modules
import prompts.executive_summary_prompts as executive_summary_prompts
import prompts.model_description_prompts as model_description_prompts
import prompts.summary_score_prompts as summary_score_prompts
import prompts.qagolden_prompts as qagolden_prompts
import nest_asyncio
from concurrent.futures import ThreadPoolExecutor, as_completed
import traceback
import httpx
import asyncio
nest_asyncio.apply()

# Initialize the OpenAI LLM model
llm = ChatOpenAI(
    model=config.openai_model,  # Change back to "gpt-4o" if that is your intended model
    openai_api_key=st.secrets["openai"]["api_key"],
    temperature=0.2,
)

# -------------------------------
# 1. Define Utility Functions
# -------------------------------

def get_model_name(df):
    latest_date = pd.to_datetime(df[df.TrainTest == "Train"].Week_Date.max())
    next_month = latest_date + pd.DateOffset(months=1)
    return f"Demand_Forecasting_Model_{next_month.strftime('%b%Y')}"


def summary_postprocessing(s):
    # Replace right single quotation mark with standard apostrophe
    s = s.replace('’', "'")
    
    # # Find the position of the first '#'
    # position = s.find('#')
    
    # # If '#' is found, return the substring after it; otherwise, return the original string
    # s = s[position:] if position != -1 else s
    
    # Replace '$' with '\$'
    return s.replace('$', '\\$')


def read_text_from_docx(file_path):
    """
    Reads and returns text from a .docx file.

    Args:
        file_path (str): The path to the .docx file.

    Returns:
        str: The extracted text from the document.
    """
    doc = docx_document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def save_session_state_to_docx(session_state_variable, file_path="executive_summary.docx"):
    """
    Saves the given session state variable content to a .docx file.

    Args:
        session_state_variable (str): The content to save.
        file_path (str): The file path including the folder and filename.
    """
    # Ensure the directory exists
    folder = os.path.dirname(file_path)
    if folder and not os.path.exists(folder):
        os.makedirs(folder)

    # Save the content to a .docx file
    doc = docx_document()
    doc.add_paragraph(session_state_variable)
    doc.save(file_path)

def create_pdf(report):
    """
    Generates a PDF from a single report dictionary with format {"Title": Content}.
    Returns the PDF as bytes.
    """
    # Extract title and content from the dictionary
    title, content = next(iter(report.items()))
    
    content = content.replace('\\$', '$')

    # Convert Markdown content to HTML
    html_content = markdown2.markdown(content, extras=["break-on-newline", "tables"])

    # Create complete HTML with improved CSS
    html = f"""
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            body {{ 
                font-family: Arial, sans-serif; 
                margin: 0.5in;  /* Reduced padding */
                line-height: 1.5; 
                font-size: 10pt; /* Increased font size */
            }}
            h1 {{ 
                text-align: center; 
                margin-bottom: 15px; 
                font-size: 12pt; /* Increased title font size */
            }}
            p, ul, li {{ 
                margin: 5px 0; 
                padding: 0; 
            }}
            ul {{ 
                padding-left: 20px; 
            }}
            li {{ 
                margin-bottom: 5px; 
            }}
        </style>
    </head>
    <body>
        <h1>{title}</h1>
        {html_content}
    </body>
    </html>
    """

    # Convert HTML to PDF
    pdf_buffer = BytesIO()
    pisa_status = pisa.CreatePDF(BytesIO(html.encode('utf-8')), dest=pdf_buffer)

    if pisa_status.err:
        raise Exception("Error converting HTML to PDF")

    return pdf_buffer.getvalue()

# Display report helper with expanders and floating download button
def display_report(session_key, report_title, file_name, height=430):
    content = st.session_state.get(session_key, "No content available.")
    # Split into sections by '## ' headings
    if "Trend" in file_name or "MultiModel" in file_name:
        pattern = r"^###\s+(.*)"
    else:
        pattern = r"^##\s+(.*)"
    lines = content.splitlines()
    sections, prelude = [], []
    current_heading, buf = None, []
    for line in lines:
        m = re.match(pattern, line)
        if m:
            if current_heading:
                sections.append((current_heading, "\n".join(buf)))
            current_heading = m.group(1).strip()
            buf = []
        else:
            (prelude if current_heading is None else buf).append(line)
    if current_heading:
        sections.append((current_heading, "\n".join(buf)))

    # Render each section in an expander
    for heading, body in sections:
        with st.expander(heading, expanded=False):
            st.markdown(body)

    # Generate and attach PDF download button
    pdf_data = create_pdf({report_title: content})
    button_container = st.container()
    with button_container:
        st.download_button(
            label="📥",
            data=pdf_data,
            file_name=file_name,
            mime="application/pdf",
        )
    button_container.float("bottom: 10px; right: -920px; z-index: 999;")
    
# Helper to show model-specific report tabs
def show_model_tabs(model_key, model_names):
    model_label = model_names[model_key]
    reports = [
        ("Model Description", "model_performance_summary_m1", "Model_Description"),
        ("Trendline Analysis", "trendline_summary_m1", "Trendline_Analysis"),
        ("Residual Shapley Analysis", "residual_shapley_summary_m1", "Residual_Shapley_Analysis"),
    ]
    sub_tabs = st.tabs([title for title, _, _ in reports])
    for i, (title, summary_key, file_prefix) in enumerate(reports):
        with sub_tabs[i]:
            display_report(
                summary_key,
                f"{model_label} - {title}",
                f"{file_prefix}_{model_label}.pdf",
                height=430
            )



def apply_custom_css():
    css = """
    <style>
        /* Hide Streamlit header and menu */
        #MainMenu, header {visibility: hidden;}

        /* Adjust main content padding */
        .block-container {
            padding-top: 1rem;
            padding-bottom: 1rem;
            padding-left: 3rem;
            padding-right: 1rem;
        }

        /* Sidebar styling */
        [data-testid="stSidebar"] {
            background-color: #053057;
            color: white;
        }
        [data-testid="stSidebar"] * {
            color: white;
        }
        [data-testid="stSidebarHeader"] img {
            height: 80px;
            width: auto;
        }
        [data-testid="stSidebar"] img {
            filter: brightness(0) invert(1);
        }
        
        [data-testid="stSidebar"] {
            min-width: 220px;
            max-width: 220px;
        }
        
        .st-key-summary div{
          color: #053057;
        }

    </style>
    """
    st.markdown(css, unsafe_allow_html=True)

def trendline_preprocessing(X_train_test_shap_M2):
    # Copy the dataset
    df = X_train_test_shap_M2.copy()

    # Define columns to exclude
    exclude_columns = [col for col in df.columns if col.endswith('_shap') or col in ['TrainTest', 'pred', 'residual']]

    # Select relevant columns
    filtered_columns = [col for col in df.columns if col not in exclude_columns]
    df_subset = df[filtered_columns].copy()

    # Convert 'Week_Date' to datetime
    df_subset['Week_Date'] = pd.to_datetime(df_subset['Week_Date'])

    # Define rolling window size
    window_size = 5

    # Apply rolling mean to 'y' (target variable)
    df['rolling_mean_y'] = pd.to_numeric(df['y']).rolling(window=window_size).mean()

    # Filter out specific features for rolling mean calculation
    features_to_exclude = ['Store_ID', 'Week_Date', 'IsHoliday', 'Type_A', 'Type_B', 'Type_C', 'year', 'week_of_month', 'week_of_year']
    features = [col for col in filtered_columns if col not in features_to_exclude]

    # Apply rolling mean to selected features
    for feature in features:
        df_subset[f'rolling_mean_{feature}'] = df_subset[feature].rolling(window=window_size).mean()
    

    # Calculate Interquartile Range (IQR) for outlier detection
    Q1, Q3 = df_subset['y'].quantile([0.25, 0.75])
    IQR = Q3 - Q1
    lower_bound, upper_bound = Q1 - 1.1 * IQR, Q3 + 1.1 * IQR

    # Identify significant highs and lows
    df_subset['is_significant_high'] = df_subset['y'] > upper_bound
    df_subset['is_significant_low'] = df_subset['y'] < lower_bound

    # Extract and filter high data points
    df_subset_high = df_subset[df_subset['is_significant_high']].reset_index(drop=True)
    df_subset_high_filtered = df_subset_high.drop(columns=['is_significant_high', 'is_significant_low'])

    # Convert each row into a JSON object and display
    json_objects = df_subset_high_filtered.to_dict(orient='records')
    json_objects = df_subset.to_dict(orient='records')
    json_objects = sorted(json_objects, key=lambda x: (x['Week_Date'], x['Store_ID']))
    
    return json_objects

def forecast_discrepancy_analysis_preprocessing(X_train_test_shap_M1, X_train_test_shap_M2):
    
    # Subset to pull only Test records
    X_train_test_shap_M1_subset = X_train_test_shap_M1[X_train_test_shap_M1['TrainTest'] == 'Test']
    X_train_test_shap_M2_subset = X_train_test_shap_M2[X_train_test_shap_M2['TrainTest'] == 'Test']

    # Convert 'Week_Date' to datetime for both datasets
    X_train_test_shap_M1_subset['Week_Date'] = pd.to_datetime(X_train_test_shap_M1_subset['Week_Date'])
    X_train_test_shap_M2_subset['Week_Date'] = pd.to_datetime(X_train_test_shap_M2_subset['Week_Date'])


    # Common dates between both datasets
    common_dates = X_train_test_shap_M1_subset['Week_Date'].isin(X_train_test_shap_M2_subset['Week_Date'])

    # Filter both datasets for the common weeks
    M1_common_period = X_train_test_shap_M1_subset[common_dates]
    M2_common_period = X_train_test_shap_M2_subset[X_train_test_shap_M2_subset['Week_Date'].isin(M1_common_period['Week_Date'])]

    # Merge the datasets on 'Week_Date' and 'store_id'
    merged_df = pd.merge(M1_common_period, M2_common_period, on=['Week_Date', 'Store_ID'], suffixes=('_model_1', '_model_2'))

    # Calculate the forecast differences
    merged_df['forecast_diff'] = abs(merged_df['pred_model_1'] - merged_df['pred_model_2'])

    # IQR method works well when you have a data distribution that's skewed or has extreme values. It’s robust against non-normal distributions.

    # IQR Outlier Detection
    Q1 = merged_df['forecast_diff'].quantile(0.25)
    Q3 = merged_df['forecast_diff'].quantile(0.75)
    IQR = Q3 - Q1

    # Define the outlier threshold
    upper_bound = Q3 + 1.5 * IQR

    # Identify outliers based on the IQR method
    merged_df['is_outlier_IQR'] = (merged_df['forecast_diff'] > upper_bound)

    # Get the outliers
    outliers_IQR = merged_df[merged_df['is_outlier_IQR']].reset_index(drop=True)

    # Convert each row of the DataFrame into a JSON object and store in a list
    json_objects = [row.to_dict() for _, row in outliers_IQR.iterrows()]

    return json_objects


def subset_df_w_high_residuals(X_train_test_shap_M1_subset_copy):
  # Calculate the IQR for residuals of both models
  X_train_test_shap_M1_subset_copy['abs_residual'] = abs(X_train_test_shap_M1_subset_copy['residual'])

  Q1_m1 = X_train_test_shap_M1_subset_copy['abs_residual'].quantile(0.25)
  Q3_m1 = X_train_test_shap_M1_subset_copy['abs_residual'].quantile(0.75)

  IQR_m1 = Q3_m1 - Q1_m1
  upper_bound_m1 = Q3_m1 + 1.5 * IQR_m1
  X_train_test_shap_M1_subset_copy['outlier_residual'] = (X_train_test_shap_M1_subset_copy['abs_residual'] > upper_bound_m1)

  residual_IQR_m1 = X_train_test_shap_M1_subset_copy[X_train_test_shap_M1_subset_copy['outlier_residual']].reset_index(drop=True)
  return residual_IQR_m1


def residual_shapley_preprocessing(X_train_test_shap_M1):
  #Make a local copy
  X_train_test_shap_M1_subset = X_train_test_shap_M1[X_train_test_shap_M1['TrainTest'] == 'Test']
  X_train_test_shap_M1_subset_copy = X_train_test_shap_M1_subset.copy()

  outlier_residual_m1 = subset_df_w_high_residuals(X_train_test_shap_M1_subset_copy)

  # Pull all column names with "_model_1" suffix
  model_1_columns = [col for col in outlier_residual_m1.columns if col.endswith('_shap') or col in ['Store_ID', 'Week_Date', 'residual', 'y', 'pred']]
  # Exclude columns based on the list of names
  exclude_columns = ["TrainTest_model_1"]
  filtered_columns = [col for col in model_1_columns if col not in exclude_columns]

  #Subset the data with only the specific shapley and pred/residual columns
  outlier_residual_m1_shap = outlier_residual_m1[filtered_columns]

  # Convert each row of the DataFrame into a JSON object and store in a list
  json_objects1 = [row.to_dict() for _, row in outlier_residual_m1_shap.iterrows()]

  # Step 1: Extract the columns that contain Shapley values (e.g., columns ending with "_shap")
  shap_columns = [col for col in X_train_test_shap_M1_subset_copy.columns if col.endswith('_shap')]
  avg_shap_values = X_train_test_shap_M1_subset_copy[shap_columns].mean()
  avg_shap_json_m1 = avg_shap_values.to_dict()
  
  return json_objects1, avg_shap_json_m1

# ------------------------- MapReduce Functions -------------------------
def map_summary(map_prompt_template: str, row_data: dict, data_dict: str, avg_shapley_values:str, pred_var:str, forecast_disc:str) -> str:
    """Refines an individual analysis summary using the detailed MAP prompt."""
    prompt = PromptTemplate(input_variables=["row_data", "data_dict", "avg_shapley_values", "pred_var", "forecast_disc"], template=map_prompt_template)
    chain = LLMChain(llm=llm, prompt=prompt, output_parser=StrOutputParser())
    refined = chain.run({"row_data": row_data, "data_dict": data_dict, "avg_shapley_values":avg_shapley_values, "pred_var": pred_var, "forecast_disc":forecast_disc})
    st.write("Map_summary_output:",refined.strip())
    return refined.strip()

def reduce_two_summaries(reduce_prompt_template: str, summary1: str, summary2: str, gold_template: str, data_dict: str, key_category: str, pred_var: str, forecast_disc:str) -> str:
    """Merges two refined summaries into a unified aggregated report using the detailed REDUCE prompt."""
    prompt = PromptTemplate(
        input_variables=["summary1", "summary2", "gold_template", "data_dict", "key_category", "pred_var", "forecast_disc"],
        template=reduce_prompt_template
    )
    chain = LLMChain(llm=llm, prompt=prompt, output_parser=StrOutputParser())
    combined = chain.run({
        "summary1": summary1,
        "summary2": summary2,
        "gold_template": gold_template,
        "data_dict": data_dict,
        "key_category": key_category,
        "pred_var": pred_var,
        "forecast_disc":forecast_disc
    })
    return combined.strip()

def append_json(file_path, new_data):
    # Step 1: Load existing data if file exists
    if os.path.exists(file_path):
        with open(file_path, "r", encoding="utf-8") as file:
            try:
                existing_data = json.load(file)
            except json.JSONDecodeError:
                existing_data = []
    else:
        existing_data = []

    # Step 2: Ensure data is in list format for appending
    if not isinstance(existing_data, list):
        existing_data = [existing_data]
    if isinstance(new_data, list):
        existing_data.extend(new_data)
    else:
        existing_data.append(new_data)

    # Step 3: Write the updated data back to file
    with open(file_path, "w", encoding="utf-8") as file:
        json.dump(existing_data, file, indent=4, ensure_ascii=False)


def map_reduce(json_objects,
               retail_data_dict,
               gold_template,
               map_prompt_template,
               reduce_prompt_template,
               key_category,
               target_variable,
               tag,
               avg_shapley_values="",
               forecast_disc=""):

    # 1) MAP PHASE
    tag_dict = {
        "residual_shapley_m1": "Shapley",
        "residual_shapley_m2": "Shapley",
        "forecast_discrepancy": "Forecast Discrepancy",
        "trendline_m1": "Trendline|Anomaly|Outlier",
        "trendline_m2": "Trendline|Anomaly|Outlier",
    }

    # parallel map
    with ThreadPoolExecutor() as executor:
        summaries = list(executor.map(
            lambda obj: map_summary(
                map_prompt_template,
                obj,
                retail_data_dict,
                avg_shapley_values,
                target_variable,
                forecast_disc
            ),
            json_objects
        ))
    

    # annotate and overwrite JSON once
    folder = "json_files"
    os.makedirs(folder, exist_ok=True)
    file_path = os.path.join(folder, f"{tag}.json")

    json_data = [
        {
            "summary": summary,
            "Week_Date": str(obj["Week_Date"]),
            "Store_ID": obj["Store_ID"],
            "Tag":     tag_dict[tag],
        }
        for obj, summary in zip(json_objects, summaries)
    ]

    # overwrite the JSON file
    with open(file_path, "w", encoding="utf-8") as f:
        json.dump(json_data, f, indent=4, ensure_ascii=False)

    # 2) REDUCE PHASE
    def reduce_list(items):
        
        while len(items) > 1:
            # next_round = []
            new_summaries = []
            with ThreadPoolExecutor() as exec2:
                futures = []
                for i in range(0, len(items) - 1, 2):
                    futures.append(exec2.submit(
                        reduce_two_summaries,
                        reduce_prompt_template,
                        items[i],
                        items[i+1],
                        gold_template,
                        retail_data_dict,
                        key_category,
                        target_variable,
                        forecast_disc
                    ))
                if len(items) % 2 == 1:
                    new_summaries.append(items[-1])
                for future in as_completed(futures):
                    new_summaries.append(future.result())
            items = new_summaries
            
        return items[0]

    # split summaries in two
    mid = len(summaries) // 2
    s1 = reduce_list(summaries[:mid])
    s2 = reduce_list(summaries[mid:])
    # st.write("summaries:\n\n\n",summaries)
    # return summaries[0]

    # FINAL MERGE
    return reduce_two_summaries(
        reduce_prompt_template,
        s1, s2,
        gold_template,
        retail_data_dict,
        key_category,
        target_variable,
        forecast_disc
    )
    



#---------------------------------------------------------------------------------------------------------------
# Model Description
#---------------------------------------------------------------------------------------------------------------
def calculate_metrics(actual, predicted):
    """
    Calculate RMSE, MAE, and SMAPE for given actual and predicted arrays.
    """
    rmse = np.sqrt(np.mean((actual - predicted) ** 2))
    mae = np.mean(np.abs(actual - predicted))
    smape = 100 * np.mean(2 * np.abs(actual - predicted) / (np.abs(actual) + np.abs(predicted)))
    return rmse, mae, smape

def get_high_correlations(X_train_test_shap_M1, threshold=0.8):
    """
    Extracts high correlations from the correlation matrix for features without '_shap' 
    in their names and without 'Store_ID', 'Week_Date', 'TrainTest', 'residual'.
    """
    # Exclude specified columns and SHAP-related features
    excluded_columns = {'Store_ID', 'Week_Date', 'TrainTest', 'residual'}
    feature_columns = [col for col in X_train_test_shap_M1.columns 
                       if '_shap' not in col and col not in excluded_columns]

    # Compute correlation matrix
    correlation_matrix = X_train_test_shap_M1[feature_columns].corr().abs()

    # Extract high correlations
    high_corr_pairs = []
    for i in range(len(feature_columns)):
        for j in range(i + 1, len(feature_columns)):  # Avoid duplicates
            if correlation_matrix.iloc[i, j] >= threshold:
                high_corr_pairs.append((feature_columns[i], feature_columns[j], correlation_matrix.iloc[i, j]))

    # Convert to DataFrame
    high_corr_df = pd.DataFrame(high_corr_pairs, columns=['Feature_1', 'Feature_2', 'Correlation'])

    return high_corr_df

def generate_total_model_summary(business_scenario, model_info, model_hyperparams,
                                 metrics_train, metrics_test, feature_importance,
                                 p_value, dw_stat, column_description, high_corr_df, target_data_desc, corr_thresh=0.8):
    """
    Generates a structured business-oriented model summary.
    """
    prompt = model_description_prompts.generate_prompt(business_scenario, model_info, model_hyperparams,
                             metrics_train, metrics_test, feature_importance,
                             p_value, dw_stat, column_description, high_corr_df, target_data_desc, corr_thresh)
    
    client = openai.OpenAI(api_key=st.secrets["openai"]["api_key"])
    chat_completion = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a helpful AI assistant."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2
    )
    
    return chat_completion.choices[0].message.content


def model_description_preprocessing(X_train_test_shap_M1):
    """
    Processes test data for residual analysis, calculates feature importance,
    performs statistical tests, and computes model performance metrics.
    """
    # high correlations
    high_corr_df = get_high_correlations(X_train_test_shap_M1)

    # Separate train and test data
    df_train = X_train_test_shap_M1[X_train_test_shap_M1['TrainTest'] == 'Train']
    df_test = X_train_test_shap_M1[X_train_test_shap_M1['TrainTest'] == 'Test']

    # Extract residuals from test data
    test_residual = df_test['residual']

    # Calculate feature importance using SHAP values
    shap_columns = [col for col in X_train_test_shap_M1.columns if 'shap' in col]
    feature_importance = X_train_test_shap_M1[shap_columns].abs().mean().sort_values(ascending=False)

    # Perform Breusch-Pagan test for heteroscedasticity
    X_train = df_test.drop(columns=['residual', 'Week_Date'])
    X_train = pd.get_dummies(X_train, drop_first=True).fillna(0)  # Convert categorical to dummy variables
    X_train_const = sm.add_constant(X_train)
    _, p_value, _, _ = het_breuschpagan(test_residual, X_train_const)

    # Perform Durbin-Watson test for autocorrelation
    dw_stat = durbin_watson(test_residual)

    # Extract actual and predicted values for both train and test sets
    actual_train = df_train['y']
    predicted_train = df_train['pred']
    actual_test = df_test['y']
    predicted_test = df_test['pred']

    # Compute performance metrics
    train_metrics = calculate_metrics(actual_train, predicted_train)
    test_metrics = calculate_metrics(actual_test, predicted_test)

    # Organize results into a dictionary
    results = {
        "breusch_pagan_p_value": round(p_value, 4),
        "durbin_watson_statistic": round(dw_stat, 2),
        "feature_importance": feature_importance,
        "train_metrics": {"rmse": train_metrics[0], "mae": train_metrics[1], "smape": train_metrics[2]},
        "test_metrics": {"rmse": test_metrics[0], "mae": test_metrics[1], "smape": test_metrics[2]},
        "high_corr_df": high_corr_df,
        "target_data_desc":X_train_test_shap_M1.y.describe()
    }
    return results

#--------------------------------------------------------------------------------
# executive Summary
#--------------------------------------------------------------------------------

def generate_executive_summary(gold_template, business_scenario, data_dict, model_description_summary,
                    trendline_analysis_summary, residual_analysis_summary, forecast_analysis_summary = None):
    """
    Generates a structured business-oriented model summary.
    """
    prompt = executive_summary_prompts.generate_prompt(gold_template, business_scenario, data_dict, model_description_summary,
                    forecast_analysis_summary, trendline_analysis_summary, residual_analysis_summary)
    client = openai.OpenAI(api_key=st.secrets["openai"]["api_key"])
    chat_completion = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a helpful AI assistant."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2
    )
    
    return chat_completion.choices[0].message.content

#-------------------------------------------------------------------------------------------------
# Chat with your Data functions
#-------------------------------------------------------------------------------------------------
def display_filter(df):
    st.markdown("<br><br><br>", unsafe_allow_html=True)
    fc = st.container()
    with fc:
        st.title("💬 Chat with your Data")
        
        col1, col2 = st.columns([2, 3])
        if 'df_filtered' not in st.session_state:
            st.session_state.df_filtered = df  # Default: no filtering applied
        if 'clear_filters' not in st.session_state:
            st.session_state.clear_filters = False 
        
        with col2:
            # Get shapes for display
            rows, cols = df.shape
        
            if 'filter_name' not in st.session_state:
                st.session_state.filter_name = f"Data Filters:({rows}, {cols})"
                
            with st.popover(st.session_state.filter_name, use_container_width=True):
                # Here you can add additional UI elements for setting filter parameters
                temp = filter_dataframe(df, st.session_state.clear_filters)
                st.session_state.clear_filters = False
                cola, colb, _ = st.columns([1,1,2])
                with cola:
                    if st.button("Apply Filter"):
                        st.session_state.df_filtered = temp
                        rows_f, cols_f = st.session_state.df_filtered.shape
                        st.session_state.filter_name = f"Data Filters:({rows}, {cols}) → ({rows_f}, {cols_f})"
                        st.rerun()
                with colb:
                    if st.button("Clear Filters"): 
                        st.session_state.clear_filters = True
                        st.session_state.df_filtered = df
                        st.session_state.filter_name = f"Data Filters:({rows}, {cols})"
                        st.rerun()
        
        with col1:            
                # Initialize sample questions
            if "sample_questions" not in st.session_state:
                st.session_state.sample_questions = [
                    "Select an option",
                    "What is the MAPE of the model across the regions?",
                    "How does the error (residual) distribution look? Is it biased in any way? Generate a plot and explain what it means",
                    "How does the actual(y) vs prediction(pred) look? Do the model predictions close to actuals? Generate a plot and explain what it means",
                    "What is overall actual(y) and prediction(pred) revenue across different months and quarters?"
                    "Are there particular months where the model underperforms?",
                    "How does the model performance vary by region (DSX, BTS)?",
                    "Is the model systematically overestimating or underestimating sales during different quarter?",
                    "What are the top 5 most influential features according to SHAP values?",
                    "How do external factors (e.g.,Unemployment, FedFund, CPI) impact revenue in the data?",
                    "Are certain features disproportionately influencing the predictions for specific stores?",
                    # "Are the SHAP values for different lagged sales features (weekly_sales_lag_1w, weekly_sales_lag_52w) in line with expectations?",
                    "How does the model capture weekly and yearly trends effectively?",
                    "How well does it predict sales for different seasons (e.g., summer vs. winter)?",
                    "Does the model perform differently across different regions, month_of_year and quarter_of_year?",
                    # "How does the model perform across different store types (Type_A, Type_B, Type_C)?",
                    # "Do stores of similar sizes (Size_shap) exhibit similar prediction patterns?",
                    # "How does the model handle sales fluctuations around holidays (IsHoliday)?",
                    # "How much does recent sales history (weekly_sales_lag_1w, weekly_sales_lag_52w, weekly_sales_lag_104w) influence predictions?",
                    # "Are there cases where the model is overly reliant on lagged sales rather than external factors?"
                ]

                # Display sample questions dropdown and button in a form
            
            user_input = "Select an option"
            if st.session_state.sample_questions:
                with st.popover("Sample Questions" , use_container_width=True):
                    user_input = st.selectbox(
                        "Choose a sample question:",
                        st.session_state.sample_questions,
                        key="sample_questions_select"
                    )
                    

    return st.session_state.df_filtered, user_input

def filter_dataframe(df: pd.DataFrame, clear_filters) -> pd.DataFrame:
    """
    Adds a UI on top of a dataframe to let viewers filter columns

    Args:
        df (pd.DataFrame): Original dataframe

    Returns:
        pd.DataFrame: Filtered dataframe
    """
        
    df = df.copy()
    # Try to convert datetimes into a standard format (datetime, no timezone)
    for col in df.columns:
        if is_object_dtype(df[col]):
            try:
                df[col] = pd.to_datetime(df[col])
            except Exception:
                pass

        if is_datetime64_any_dtype(df[col]):
            df[col] = df[col].dt.tz_localize(None)
        
    modification_container = st.container()
    
    if 'to_filter_columns' not in st.session_state:
        st.session_state.to_filter_columns = []
        
    if clear_filters:
        st.session_state.to_filter_columns = []

    with modification_container:
        temp = st.multiselect("Filter dataframe on", df.columns, default=st.session_state.to_filter_columns)
        if temp != st.session_state.to_filter_columns:
            st.session_state.to_filter_columns = temp
            st.rerun()
        for column in temp:
            # Treat columns with < 10 unique values as categorical
            if is_categorical_dtype(df[column]) or df[column].nunique() < 10:
                left, right = st.columns((1, 20))
                left.write("↳")
                user_cat_input = right.multiselect(
                    f"Values for {column}",
                    df[column].unique(),
                    default=list(df[column].unique()),
                )
                df = df[df[column].isin(user_cat_input)]
            elif is_numeric_dtype(df[column]):
                # Compute the min and max for the column
                left, col1, col2, col3 = st.columns([1,7,7,7])
                left.write("↳")
                _min = float(df[column].min())
                _max = float(df[column].max())
                # step = (_max - _min) / 100
                
                if pd.api.types.is_integer_dtype(df[column]):
                    step = 1.0
                else:
                    step = (_max - _min) / 100

                with col1:
                    lower_bound = st.number_input(
                        "Min", value=_min, min_value=_min, max_value=_max, step=step
                    )

                with col2:
                    st.markdown("######")
                    st.markdown(f"""&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;<=&nbsp;&nbsp;&nbsp;{column}&nbsp;&nbsp;&nbsp;<=""", unsafe_allow_html=True)

                with col3:
                    upper_bound = st.number_input(
                        "Max", value=_max, min_value=_min, max_value=_max, step=step
                    )

                # Optionally, ensure lower_bound is not greater than upper_bound
                if lower_bound > upper_bound:
                    right.error("The minimum value cannot be greater than the maximum value.")

                # Filter the dataframe using the selected bounds
                df = df[(df[column] >= lower_bound) & (df[column] <= upper_bound)]
            elif is_datetime64_any_dtype(df[column]):
                left, right = st.columns((1, 20))
                left.write("↳")
                user_date_input = right.date_input(
                    f"Values for {column}",
                    value=(
                        df[column].min(),
                        df[column].max(),
                    ),
                )
                
                if len(user_date_input) == 2:
                    user_date_input = tuple(map(pd.to_datetime, user_date_input))
                    start_date, end_date = user_date_input
                    df = df.loc[df[column].between(start_date, end_date)]
            else:
                left, right = st.columns((1, 20))
                left.write("↳")
                user_text_input = right.text_input(
                    f"Substring or regex in {column}",
                )
                if user_text_input:
                    df = df[df[column].astype(str).str.contains(user_text_input)]
    return df



def clear_saved_charts():
    """Clears the saved_charts folder if there are no messages in session state."""
    folder = "saved_charts"
    if "messages" in st.session_state and not st.session_state.messages:
        for filename in os.listdir(folder):
            file_path = os.path.join(folder, filename)
            try:
                if os.path.isfile(file_path):
                    os.remove(file_path)
            except Exception as e:
                st.error(f"Error deleting file {file_path}: {e}")

def save_chart_with_unique_name(source_path, destination_folder):
    """Saves the chart with a unique filename based on the timestamp."""
    if os.path.exists(source_path):
        os.makedirs(destination_folder, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        destination_path = os.path.join(destination_folder, f"chat_chart_{timestamp}.png")
        shutil.move(source_path, destination_path)
        return destination_path
    return None

def chatbot(df, sample_q):
    """Main chatbot function handling messages and chart generation."""

    # Initialize chat history
    if "messages" not in st.session_state:
        st.session_state.messages = []

    # Clear saved charts if chat history is empty
    clear_saved_charts()

    # Display chat history
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            if message.get("type") == "image":
                st.image(message["content"], caption="Generated Chart")
            else:
                st.markdown(str(message["content"]).replace('$', '\\$'))

    # Get user input from chat
    user_input = st.chat_input("Type your message here...")
    if sample_q != "Select an option" and user_input is None:
        user_input = sample_q

    # Process user input (from chat or sample question)
    if user_input and user_input != "Select an option":
        # Append user message to history
        st.session_state.messages.append({"role": "user", "content": user_input})

        # Display user message
        with st.chat_message("user"):
            st.markdown(user_input)

        # Get response from AI
        with st.spinner("Thinking..."):
            df1 = pai.DataFrame(df)
            df2 = pai.DataFrame(st.session_state.get("retail_data_dictionary", {}))  # Handle missing dictionary
            
            max_attempts = 10
            attempt = 0
            success = False

            original_input = user_input  # Save original input
            while attempt < max_attempts and not success:
                try:
                    if attempt == 0:
                        # First attempt, use original input
                        assistant_reply = pai.chat(user_input, df1, df2)
                    else:
                        # Subsequent attempts, add the clarification
                        modified_input = original_input + ", Only asked for a plot then save it with location as a raw string"
                        assistant_reply = pai.chat(modified_input, df1, df2)
                    success = True
                except Exception as e:
                    st.write(f"Attempt {attempt + 1} failed with error: {e}")
                    attempt += 1


            if not success:
                st.write("All 10 attempts failed.")
                
            import base64
            assistant_reply_str = str(assistant_reply)
            chart_source = "exports/charts/temp_chart.png"

            if assistant_reply_str.strip().startswith("data:image/png;base64,"):
                # Extract base64 part
                b64_string = assistant_reply_str.split(",", 1)[1].strip()
                
                # Decode and save the image
                with open(chart_source, "wb") as f:
                    f.write(base64.b64decode(b64_string))

            
            template = f"""
            Business Scenario:
            {config.business_scenario}

            User Question:
            {user_input}

            PandasAI Output:
            {assistant_reply}

            Instructions:
            You are a technical assistant supporting analytics for the given business scenario. Based on the user's question and the PandasAI output, internally perform the following:

            1. Determine the most appropriate format for the response (paragraph, table, or plot).
            2. If needed, reformat and enhance the PandasAI output to be clearer, more precise, and insightful.
            3. Ensure the output is relevant to the business scenario. If not, return:
            "I am unable to generate a response, please ask a different question."
            4. If the output contains errors, return a clear explanation and suggest a fix.

            DO NOT include any of the steps above in your response. ONLY return the final output.


            """

                
            # Set your API key
            client = openai.OpenAI(api_key=st.secrets["openai"]["api_key"])
            # Call the ChatGPT model
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": template}
                ]
            )

            # Get the reply
            assistant_reply = response.choices[0].message.content
        


        chart_destination_folder = "saved_charts"
        saved_chart_path = save_chart_with_unique_name(chart_source, chart_destination_folder)

        if saved_chart_path:
            # Append and display the generated chart
            st.session_state.messages.append({"role": "assistant", "type": "image", "content": saved_chart_path})
            with st.chat_message("assistant"):
                st.image(saved_chart_path, caption="Generated Chart")
        else:
            # Display assistant response
            with st.chat_message("assistant"):
                st.markdown(str(assistant_reply).replace('$', '\\$'))

            # Append assistant response to history
            st.session_state.messages.append({"role": "assistant", "content": assistant_reply})
            
#--------------------------------------------------------------------------
# Chat with reports
#--------------------------------------------------------------------------

# Function to load documents from Word and JSON with consistent metadata handling
def load_documents(data_files, json_file_data):
    """Load both Word and JSON documents with unified metadata handling"""
    docs = []

    # Load Word documents
    for file in data_files:
        try:
            from langchain.document_loaders import UnstructuredWordDocumentLoader
            loader = UnstructuredWordDocumentLoader(file)
            word_docs = loader.load()
            for doc in word_docs:
                # Extract file name without extension for easier matching
                file_name = file.split('.')[0]

                metadata = {
                    'source_type': 'word',
                    'filename': file,
                    'document_type': file_name,
                    'needs_splitting': True  # Flag to indicate this document needs splitting
                }

                file_name = file.split('.')[0].lower()
                if 'shapley' in file_name:
                    metadata['tag'] = 'shapley'
                elif 'forecast discrepancy' in file_name:
                    metadata['tag'] = 'forecast_discrepancy'
                elif 'trendline' in file_name:
                    metadata['tag'] = 'trendline'
                else:
                    metadata['tag'] = 'unknown'
                doc.metadata = metadata
                docs.append(doc)
        except Exception as e:
            print(f"Error loading Word document {file}: {e}")

    # Load JSON documents with date formatting
    for json_file in json_file_data:
        try:
            with open(json_file, 'r') as file:
                json_data = json.load(file)

            for item in json_data:
                metadata = {
                    'source_type': 'json_summary',
                    'needs_splitting': False  # Flag to indicate this document does NOT need splitting
                }

                # Extract year, month, day from Week_Date
                if 'Week_Date' in item and item['Week_Date']:
                    date_str = str(item['Week_Date'])

                    # Try to parse the date and extract components
                    try:
                        # Expected format: YYYY-MM-DD
                        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                        # Store as integers
                        metadata['year'] = int(date_obj.year)
                        metadata['month'] = int(date_obj.month)
                        metadata['day'] = int(date_obj.day)
                    except ValueError:
                        # If parsing fails, try to extract using regex
                        date_match = re.search(r'(\d{4})-(\d{1,2})-(\d{1,2})', date_str)
                        if date_match:
                            metadata['year'] = int(date_match.group(1))
                            metadata['month'] = int(date_match.group(2))
                            metadata['day'] = int(date_match.group(3))
                        else:
                            print(f"Could not parse date: {date_str}")

                # Add store_id if present
                if 'Store_ID' in item and item['Store_ID']:
                    metadata['store_id'] = str(item['Store_ID'])
                if 'tag' in item and item['tag']:
                    metadata['tag'] = item['tag']

                # Add other metadata fields if present
                if 'nda' in item and item['nda'] is not None:
                    metadata['nda'] = str(item['nda'])

                doc = Document(
                    page_content=item['summary'],
                    metadata=metadata
                )

                docs.append(doc)

        except (FileNotFoundError, json.JSONDecodeError) as e:
            print(f"Error loading JSON file: {e}")

    return docs


def initialize_rag_system():
    """Initialize the RAG system with document loading and vectorization"""
    data_files = [
        "summary_reports/trendline_summary_m1.docx",
        "summary_reports/residual_shapley_summary_m1.docx"
    ]
    
    json_data_files = [
        "json_files/trendline_m1.json",
        "json_files/residual_shapley_m1.json"
    ]
    
    if st.session_state.get("m2_flag", False):
        data_files += [
            "summary_reports/trendline_summary_m2.docx",
            "summary_reports/forecast_discrepancy_summary.docx",
            "summary_reports/residual_shapley_summary_m2.docx"
        ]
        
        json_data_files += [
            "json_files/trendline_m2.json",
            "json_files/forecast_discrepancy.json",
            "json_files/residual_shapley_m2.json"
        ]


    docs = load_documents(data_files, json_data_files)

    # Initialize text splitter for documents that need splitting
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)

    # Process documents based on whether they need splitting
    final_documents = []

    for doc in docs:
        if doc.metadata.get('needs_splitting', False):
            # Split this document
            split_docs = text_splitter.split_documents([doc])
            final_documents.extend(split_docs)
        else:
            # Keep this document as is (don't split)
            final_documents.append(doc)

        # Remove the temporary 'needs_splitting' flag
        if 'needs_splitting' in doc.metadata:
            del doc.metadata['needs_splitting']

    # Ensure all metadata values are simple strings, integers, floats, or booleans
    for doc in final_documents:
        for key, value in list(doc.metadata.items()):
            if not isinstance(value, (str, int, float, bool)) or value is None:
                if value is None:
                    doc.metadata[key] = ""
                else:
                    doc.metadata[key] = str(value)
        # Check content-metadata consistency
        if 'store_id' in doc.metadata and doc.metadata['store_id']:
            if doc.metadata['store_id'] not in doc.page_content:
                print(f"Warning: store_id {doc.metadata['store_id']} not in content: {doc.page_content[:100]}")

    # Create embeddings and vector store
    embeddings = OpenAIEmbeddings(
        model="text-embedding-3-large",
        dimensions=1024,
        openai_api_key=st.secrets["openai"]["api_key"]
    )

    db = Chroma.from_documents(final_documents, embeddings, persist_directory="./chroma_db")

    return 


def format_docs(docs):
    """Format documents with their metadata for better context"""
    formatted_docs = []

    for i, doc in enumerate(docs):
        metadata_str = ", ".join([f"{k}: {v}" for k, v in doc.metadata.items()
                               if k in ['tag','store_id', 'year', 'month', 'day', 'document_type', 'source_type']])

        formatted_doc = f"Metadata: {metadata_str}\n\nDocument {i+1}:\n{doc.page_content}\n{'-'*50}"
        formatted_docs.append(formatted_doc)

    return "\n\n".join(formatted_docs)


# Define query rewriter globally (or pass llm explicitly)
query_rewrite_prompt = PromptTemplate(
    input_variables=["question"],
    template="""
    Rewrite the following question to clarify its intent for a retail analytics search system. Separate the core query from any metadata hints (e.g., store ID, date) without assuming specific formats. Output as JSON with 'query' and 'metadata_hints' fields.

    Question: {question}

    Example:
    Input: "Why did Store ID 10 see a sales spike on November 28, 2010?"
    Output: {{"query": "why sales spike", "metadata_hints": {{"store": "10", "date": "November 28, 2010"}}}}
    Input: "What combination of factors led to Store 14's exceptional performance in the last week of November 2010?"
    Output: {{"query": "Store 14 November 2010 exceptional performance", "metadata_hints": {{"store_id": "14", "year": "2010", "month": "11"}}}}
    Output:
    """
)

def create_ensemble_retriever(db, llm):
    # Base retrievers from vector store
    similarity_retriever = db.as_retriever(search_type="similarity", search_kwargs={"k": 10})
    mmr_retriever = db.as_retriever(search_type="mmr", search_kwargs={"k": 10, "fetch_k": 20, "lambda_mult": 0.5})

    # Metadata retriever setup
    metadata_field_info = [
        AttributeInfo(name="store_id", description="Exact store ID (e.g., '10'). Optional in queries.", type="string"),
        AttributeInfo(name="year", description="Year of the document date. Use with month and day if date is specified.", type="integer"),
        AttributeInfo(name="month", description="Month of the document date (1-12). Use with year and day.", type="integer"),
        AttributeInfo(name="day", description="Day of the document date (1-31). Use with year and month.", type="integer"),

    ]

    document_content_description = """
    Retail analytics documents with sales anomalies, forecasts, and Shapley analyses. The system MUST:
    1. Match store_id exactly if specified (e.g., 'Store 10' means store_id='10').
    2. Match year, month, and day together if a specific date is specified (e.g., 'November 28, 2010' means year=2010, month=11, day=28).
    3. For date ranges (e.g., 'last week of November 2010'), match year and month, omitting day if not specified.
    4. Handle queries with only dates (no store_id) by filtering on date fields only.
    5. Match report types (forecast discrepancy, shapley, trendline) based on the analytical focus of the query, - Forecast discrepancy reports the forecast discrepancies between two demand forecasting models, Model 1 and Model 2, for retail sales data.- Shapley value analysis provides a detailed analysis of forecast errors in retail sales predictions using Shapley values to quantify feature-level contributions. - Trendline analysis with anomaly detection (sales anomalies) for store performance
    6. Prioritize documents matching ALL specified metadata criteria.
    """

    examples = [
        (
            "What explains the significant 105.4% increase in Store ID 4's sales on November 28, 2010?",
            {"query": "significant 105.4% increase in sales", "filter": 'and(eq("store_id", "4"), eq("year", 2010), eq("month", 11), eq("day", 28))'}
        ),
        (
            "Why did Store ID 10 experience $2,939,946.38 in sales compared to its rolling average of $1,705,814.03 on November 28, 2010?",
            {"query": "sales compared to rolling average", "filter": 'and(eq("store_id", "10"), eq("year", 2010), eq("month", 11), eq("day", 28))'}
        ),
        (
            "How did Store ID 13's larger size contribute to its anomalous performance on November 28, 2010?",
            {"query": "larger size anomalous performance", "filter": 'and(eq("store_id", "13"), eq("year", 2010), eq("month", 11), eq("day", 28))'}
        ),
        (
            "What combination of factors led to Store 14's exceptional performance in the last week of November 2010?",
            {"query": "combination of factors exceptional performance", "filter": 'and(eq("store_id", "14"), eq("year", 2010), eq("month", 11))'}
        ),
        (
            "What trends were observed on November 28, 2010?",
            {"query": "trends observed", "filter": 'and(eq("year", 2010), eq("month", 11), eq("day", 28))'}
        ),
        (
            "Why did Store 13 show a 50.2% increase while similar stores experienced different outcomes on November 27, 2011?",
            {"query": "Store 13 November 27 2011 50.2% increase", "filter": 'and(eq("store_id", "13"), eq("year", 2011), eq("month", 11), eq("day", 27))'}
        ),
        (
            "What combination of factors led to Store 20's exceptional performance in the last week of November 2011?",
            {"query": "Store 20 November 2011 exceptional performance", "filter": 'and(eq("store_id", "20"), eq("year", 2011), eq("month", 11))'}
        ),
        ("What explains the unusual pattern where Store 13 had 68.92% higher sales than expected despite a lower unemployment rate on November 28, 2010?",
         {"query": "Store 13 November 28 2010 68.92% higher sales", "filter": 'and(eq("store_id", "13"), eq("year", 2010), eq("month", 11), eq("day", 28))'})

    ]

    constructor_prompt = get_query_constructor_prompt(
        document_content_description,
        metadata_field_info,
        allowed_comparators=["$eq", "$and"],
        examples=examples,
    )

    output_parser = StructuredQueryOutputParser.from_components()
    query_constructor = constructor_prompt | llm | output_parser

    class DebugRunnable(RunnablePassthrough):
        def invoke(self, input, config=None):
            result = super().invoke(input, config)
            print(f"Constructed filter for '{input}': {result.filter}")
            return result

    debug_query_constructor = query_constructor | DebugRunnable()

    metadata_retriever = SelfQueryRetriever(
        query_constructor=query_constructor,
        vectorstore=db,
        structured_query_translator=ChromaTranslator(),
        search_kwargs={'k': 20},
        verbose=True,
        enable_limit=True
    )

    # Query rewriting chain
    query_rewriter = query_rewrite_prompt | llm | JsonOutputParser()

    # Wrap all retrievers with invoke
    def wrap_retriever(retriever, use_full_query=False):
        def invoke_query(query):
            rewritten = query_rewriter.invoke({"question": query})
            # Use full query for metadata retriever, core query for similarity/MMR
            search_query = query if use_full_query else rewritten["query"]
            return retriever.invoke(search_query)
        return RunnableLambda(invoke_query)

    similarity_retriever_wrapped = wrap_retriever(similarity_retriever)
    mmr_retriever_wrapped = wrap_retriever(mmr_retriever)
    metadata_retriever_wrapped = wrap_retriever(metadata_retriever, use_full_query=True)

    # Ensemble retriever with wrapped retrievers
    ensemble_retriever = EnsembleRetriever(
        retrievers=[similarity_retriever_wrapped, mmr_retriever_wrapped, metadata_retriever_wrapped],
        weights=[0.2, 0.2, 0.6]
    )
    # Post-filtering to ensure exact metadata matches
    def post_filter(query, docs):
        rewritten = query_rewriter.invoke({"question": query})
        hints = rewritten["metadata_hints"]
        filtered_docs = []
        for doc in docs:
            metadata = doc.metadata
            matches = True
            if "store_id" in hints and metadata.get("store_id") != hints["store_id"]:
                matches = False
            if "year" in hints and metadata.get("year") != int(hints["year"]):
                matches = False
            if "month" in hints and metadata.get("month") != int(hints["month"]):
                matches = False
            if "day" in hints and metadata.get("day") != int(hints["day"]):
                matches = False
            if matches:
                filtered_docs.append(doc)
        return filtered_docs[:5]  # Return up to 5 exact matches

    original_invoke = ensemble_retriever.invoke
    def filtered_invoke(query):
        docs = original_invoke(query)
        return post_filter(query, docs)
    # ensemble_retriever.invoke = filtered_invoke
    # ensemble_retriever.invoke = RunnableLambda(filtered_invoke)
    # Attach retrievers for debugging (optional)
    object.__setattr__(ensemble_retriever, "similarity_retriever", similarity_retriever_wrapped)
    object.__setattr__(ensemble_retriever, "mmr_retriever", mmr_retriever_wrapped)
    object.__setattr__(ensemble_retriever, "metadata_retriever", metadata_retriever_wrapped)

    return ensemble_retriever



def process_query(query, rag_chain_with_source):
    """Process a single query using the RAG chain."""
    try:
        response = rag_chain_with_source.invoke(query)  # Synchronous call
        qa_rag = {'question': query, 'answer': response["answer"]}
        retrieved_context = extract_contentpage(response["context"])
        return qa_rag, retrieved_context
    except Exception as e:
        print(f"Error during query execution: {str(e)}")
        print(traceback.format_exc())
        return None, None


def extract_contentpage(context_sample):
    return [doc_item.page_content for doc_item in context_sample]


# ---------------------------
# Summary Score Calculation
# ---------------------------
client = openai.OpenAI(api_key=st.secrets["openai"]["api_key"])

def get_response(prompt: str) -> str:
    # resp = client.responses.create(model="gpt-4", input=prompt, temperature=0)
    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2
    )
    
    return resp.choices[0].message.content
    # return resp.output_text.strip()

def load_json(fp: str):
    return json.load(open(fp, "r", encoding="utf-8"))

def load_docx(fp: str) -> str:
    doc = docx.Document(fp)
    return "\n".join(p.text for p in doc.paragraphs)

def evaluate(prompt_template: str, individual_summary: str, ground_truth: str = "") -> int:
    prompt = prompt_template.format(
        individual_summary=individual_summary,
        ground_truth=ground_truth
    )
    return int(get_response(prompt))

def extract_detailed_summary(docx_text: str, store_id, week_date) -> str:
    return get_response(summary_score_prompts.EXTRACTION_PROMPT.format(
        docx_text=docx_text,
        store_id=store_id,
        week_date=week_date
    ))

def extract_exec_section(executive_text: str, section_heading: str) -> str:
    return get_response(summary_score_prompts.SECTION_EXTRACTION_PROMPT.format(
        executive_text=executive_text,
        section_heading=section_heading
    ))
 
 ##### generating golden Q&A pairs
def get_prompt(summary, prompt):
    
    prompt_template = PromptTemplate(template=prompt, input_variables=["summary"])
    chain = prompt_template | llm | StrOutputParser()
    # print(prompt_template)
    response = chain.invoke({"summary": summary})
    return response.strip()    
def extract_output_section(text):
    """
    Extracts the content inside the first '['.
    """
    match = re.search(r'\[([\s\S]*)', text)  # Capture everything after the first '['
    return match.group(1).strip() if match else text
def parse_qa(text):
    """
    Parses LLM-generated structured Q&A output into a list of dictionaries.

    Args:
        text (str): The structured text output from the LLM.

    Returns:
        list: A list of dictionaries, each containing 'question' and 'answer' keys.
    """
    qa_pairs = []

    # Match questions and their corresponding multiline answers
    pattern = re.findall(
        r'Question:\s*(.*?)\s*\n\s*Answer:\s*([\s\S]*?)(?=\n\s*Question:|\n?\s*\]|\Z)',
        text
    )

    for question, answer in pattern:
        qa_pairs.append({"question": question.strip(), "answer": answer.strip()})

    return qa_pairs
def save_qa_to_json(qa_list, filename="qa_output.json"):
    """
    Saves the extracted Q&A pairs to a JSON file.

    Args:
        qa_list (list): List of dictionaries containing Q&A pairs.
        filename (str): Name of the JSON file to save.
    """
    with open(filename, "w", encoding="utf-8") as f:
        json.dump(qa_list, f, indent=4, ensure_ascii=False)



def run_prompt(summary, input_qa, prompt):
  prompt_template = PromptTemplate(template=prompt, input_variables=["document_summary", "generated_qa"])
  chain = prompt_template | llm | StrOutputParser()
  response = chain.invoke({"document_summary": summary, "generated_qa": input_qa})
  return response.strip()

def score_qa_pairs(file_path, summary_doc, score_prompt):
    with open(file_path, 'r') as file:
          json_data = json.load(file)
    list_of_qa = []
    for input_qa in json_data:
        response = run_prompt(summary_doc, input_qa, score_prompt)
        input_qa["QA_score"] = response.split(":")[-1].strip()
        list_of_qa.append(input_qa)
    # remove .json from file_path
    file_path = file_path.replace(".json", "")
    with open(f"{file_path}_score.json", 'w') as file:
        json.dump(list_of_qa, file, indent=4)
    # Load the JSON file
    with open(f"{file_path}_score.json", "r", encoding="utf-8") as f:
        data = json.load(f)
    data_qa = []
    for item in data:
        if int(item["QA_score"]) > 8:
            data_qa.append(item)
    # Convert to DataFrame
    df = pd.DataFrame(data_qa, columns=["question", "answer", "QA_score"])

    # # Save to CSV
    df.to_csv(f"{file_path}_score.json", index=False, encoding="utf-8")

    # Now read the CSV back
    csv_df = pd.read_csv(f"{file_path}_score.json", encoding="utf-8")

    # Convert DataFrame back to JSON
    json_data = csv_df.to_dict(orient="records")

    # Save back to JSON file
    # remove .json from file_path
    file_path = file_path.replace(".json", "")
    # Save the final JSON file
    final_file_path = f"{file_path}_final.json"
    with open(final_file_path, "w", encoding="utf-8") as f:
        json.dump(json_data, f, ensure_ascii=False, indent=4)
    return final_file_path
def generate_golden_qa(summary_doc, type, summary):
    if type == "shapley":
        prompt = qagolden_prompts.shapley_prompt_template
    elif type == "trendline":       
        prompt = qagolden_prompts.tredline_prompt_template
    elif type == "forecast":
        prompt = qagolden_prompts.forecast_prompt_template
    response = get_prompt(summary_doc, prompt)
    
    qa_dict = parse_qa(extract_output_section(response))
    file_path = "golden_qa_files/"+f"{summary}_golden_qa.json"
    save_qa_to_json(qa_dict, file_path)
    # score the generated Q&A pairs
    score_prompt = qagolden_prompts.score_prompt_template
    final_file_path = score_qa_pairs(file_path, summary_doc, score_prompt) 
    return final_file_path

## DeepEval metrics for golden Q&A pairs
def create_test_cases(qa_data, retrieved_contexts, responses):
    test_cases = []
    for i, (qa_item, context, response) in enumerate(zip(qa_data, retrieved_contexts, responses)):
        test_case = LLMTestCase(
            input=qa_item['question'],
            actual_output=response['answer'],
            expected_output=qa_item.get('answer', ''),
            context=context,
            retrieval_context=context
        )
        test_cases.append(test_case)
    return test_cases



def process_queries_parallel(queries, rag_chain_with_source):
    rag_response = [None] * len(queries)  # Pre-allocate with original order
    retrieved_contexts = [None] * len(queries)
    with ThreadPoolExecutor(max_workers=4) as executor:
        future_to_index = {executor.submit(process_query, query, rag_chain_with_source): i for i, query in enumerate(queries)}
        for future in as_completed(future_to_index):
            index = future_to_index[future]
            try:
                qa_rag, retrieved_context = future.result()
                if qa_rag is not None and retrieved_context is not None:
                    rag_response[index] = qa_rag
                    retrieved_contexts[index] = retrieved_context
            except Exception as e:
                print(f"Error retrieving result for query '{queries[index]}': {str(e)}")
    # Filter out None values in case of failures
    rag_response = [r for r in rag_response if r is not None]
    retrieved_contexts = [c for c in retrieved_contexts if c is not None]
    return rag_response, retrieved_contexts

def save_evaluation_results(results, output_file="rag_evaluation_results.json"):
    with open(output_file, 'w') as file:
        json.dump(results, file, indent=4)
    print(f"Evaluation results saved to {output_file}")

def build_retriever():
    """Build the database and retriever for RAG system."""
    
    llm = ChatOpenAI(model=config.openai_model, temperature=0, openai_api_key=st.secrets["openai"]["api_key"])
    # db, documents = initialize_rag_system(data_files, json_data_files)
    # Create embeddings and vector store
    embeddings = OpenAIEmbeddings(
        model="text-embedding-3-large",
        dimensions=1024,
        openai_api_key=st.secrets["openai"]["api_key"]
    )
    db = Chroma(persist_directory="./chroma_db", embedding_function=embeddings)
    # print(f"Total documents: {len(documents)}")
    # print(f"Word documents (split): {len([doc for doc in documents if doc.metadata.get('source_type') == 'word'])}")
    # print(f"JSON summary documents (unsplit): {len([doc for doc in documents if doc.metadata.get('source_type') == 'json_summary'])}")

    retriever = create_ensemble_retriever(db, llm)
    return retriever
def generate_rag_responses(question, retriever):
    """Generate RAG response for a single question."""
    template_prompt = """
You are an assistant for question-answering tasks using retrieved documents. Your goal is to generate answers that are factually faithful to the provided context.

Question:
{question}

Context:
{context}

Instructions:

1. **Use Only Verified Information**:
   - Do not infer, assume, or generalize beyond the context.
   - Only include numbers, terms, or claims that appear **explicitly** in the context.
   - Do not fabricate any statistics, effects, or explanations not clearly mentioned.

2. **Matching Store IDs and Dates**:
   - If the question specifies a store ID, ensure it matches the store(s) in the context.
   - If the question specifies a date (e.g., August 26, 2012), treat it as matching if the context references that date or the week containing it e.g., “week of August 26, 2012” or "the week of 2012-08-12".
   - If there is no store ID in the question, only date relevance is required.

3. **If the context contains matching data**:
   - Answer clearly using only that data.
   - Avoid speculative statements. If something isn’t explicitly in the context, do not say it.

4. **If the context does NOT contain relevant data**:
   - If both store ID and date are specified: respond with
     "The retrieved information doesn't contain data for [store ID] on [date]."
   - If only date is specified: respond with
     "The retrieved information doesn't contain data for [date]."

5. **Language style**:
   - Be concise and accurate.
   - Do not begin your answer by confirming the data is present or not — just answer the question without any extra text.
   - Clearly explain technical terms when used (e.g., Shapley values).

Answer:
"""

    custom_rag_prompt = ChatPromptTemplate.from_template(template_prompt)

    rag_chain_from_docs = (
        RunnablePassthrough.assign(context=(lambda x: format_docs(x["context"])))
        | custom_rag_prompt
        | llm
        | StrOutputParser()
    )
    rag_chain_with_source = RunnableParallel(
        {"context": retriever, "question": RunnablePassthrough()}
    ).assign(answer=rag_chain_from_docs)

    # Process the single question
    rag_response, retrieved_context = process_query(question, rag_chain_with_source)

    if rag_response is None or retrieved_context is None:
        print("Failed to process the question.")
        return None, None, None

    # Create qa_data for evaluation (mimicking the structure from the JSON file)
    qa_data = [{'question': question, 'answer': ''}]  # Expected answer is empty since we don't have a golden answer

    return rag_response, retrieved_context, qa_data



def generate_rag_responses_goldenqa(file_path, retriever):
    
    template_prompt = """
You are an assistant for question-answering tasks using retrieved documents. Your goal is to generate answers that are factually faithful to the provided context.

Question:
{question}

Context:
{context}

Instructions:

1. **Use Only Verified Information**:
   - Do not infer, assume, or generalize beyond the context.
   - Only include numbers, terms, or claims that appear **explicitly** in the context.
   - Do not fabricate any statistics, effects, or explanations not clearly mentioned.

2. **Matching Store IDs and Dates**:
   - If the question specifies a store ID, ensure it matches the store(s) in the context.
   - If the question specifies a date (e.g., August 26, 2012), treat it as matching if the context references that date or the week containing it e.g., “week of August 26, 2012” or "the week of 2012-08-12".
   - If there is no store ID in the question, only date relevance is required.

3. **If the context contains matching data**:
   - Answer clearly using only that data.
   - Avoid speculative statements. If something isn’t explicitly in the context, do not say it.

4. **If the context does NOT contain relevant data**:
   - If both store ID and date are specified: respond with
     "The retrieved information doesn't contain data for [store ID] on [date]."
   - If only date is specified: respond with
     "The retrieved information doesn't contain data for [date]."

5. **Language style**:
   - Be concise and accurate.
   - Do not begin your answer by confirming the data is present or not — just answer the question without any extra text.
   - Clearly explain technical terms when used (e.g., Shapley values).

Answer:
"""

    custom_rag_prompt = ChatPromptTemplate.from_template(template_prompt)

    rag_chain_from_docs = (
        RunnablePassthrough.assign(context=(lambda x: format_docs(x["context"])))
        | custom_rag_prompt
        | llm
        | StrOutputParser()
    )
    rag_chain_with_source = RunnableParallel(
        {"context": retriever, "question": RunnablePassthrough()}
    ).assign(answer=rag_chain_from_docs)

    with open(file_path, 'r') as file:
        qa_data = json.load(file)
    queries = [item['question'] for item in qa_data]

    rag_response, retrieved_contexts = process_queries_parallel(queries, rag_chain_with_source)
    # Validate alignment
    for i, (q, r, c) in enumerate(zip(queries, rag_response, retrieved_contexts)):
        if q != r['question']:
            print(f"Alignment error at index {i}: query '{q}' != response question '{r['question']}'")
    
    return rag_response, retrieved_contexts, qa_data



async def close_client(client):
    await client.aclose()

async def evaluate_test_case(test_case, metric_class, threshold, model):
    client = None
    try:
        client = httpx.AsyncClient()
        metric = metric_class(threshold=threshold, model=model)
        metric.measure(test_case)  # Synchronous call
        score = metric.score if hasattr(metric, 'score') else None
        print(f"Raw score for '{test_case.input[:30]}...': {score}")
        return {
            "question": test_case.input,
            "context": test_case.context,
            "answer": test_case.actual_output,
            "retrieval_context": test_case.retrieval_context,
            "expected_answer": test_case.expected_output,
            "score": score,
            "passed": metric.is_successful(),
            "reason": metric.reason
        }
    except Exception as e:
        error_msg = f"Error evaluating {metric_class.__name__} for '{test_case.input[:30]}...': {str(e)}"
        print(error_msg)
        return {"question": test_case.input, "error": error_msg}
    finally:
        if client is not None:
            await close_client(client)

async def evaluate_saved_responses(rag_response, retrieved_contexts, qa_data, file_path):
    
    # Validate alignment before creating test cases
    for i, (qa, resp) in enumerate(zip(qa_data, rag_response)):
        if qa['question'] != resp['question']:
            print(f"Alignment error at index {i}: qa_question '{qa['question']}' != response_question '{resp['question']}'")
    test_cases = create_test_cases(qa_data, retrieved_contexts, rag_response)

    # Define metrics before calling evaluate_rag_pipeline
    model = config.openai_model
    openai.api_key = st.secrets["openai"]["api_key"]
    os.environ["OPENAI_API_KEY"] = st.secrets["openai"]["api_key"]
    metrics = [
        AnswerRelevancyMetric(threshold=0.7, model=model),
        FaithfulnessMetric(threshold=0.5, model=model),
        ContextualPrecisionMetric(threshold=0.7, model=model),
        ContextualRecallMetric(threshold=0.7, model=model)
    ]

    print("\n\n===== EVALUATING RAG PIPELINE =====")
    eval_results = await evaluate_rag_pipeline(test_cases, metrics)
    # remove json from file_path
    file_path = file_path.replace(".json", "")
    # Save evaluation results
    save_evaluation_results(eval_results, f"{file_path}_eval_results.json")
    eval_results_avg = {}
    for metric_name, result in eval_results.items():
        eval_results_avg[metric_name] = result["average_score"]   
        
    print(f"Evaluation results: {eval_results_avg}")
    # print("Evaluation results saved to rag_evaluation_results.json")
    return eval_results_avg

async def evaluate_rag_pipeline(test_cases, metrics=None):
    results = {}
    for metric in metrics:
        metric_name = metric.__class__.__name__
        print(f"\nEvaluating {metric_name}...")
        metric_results = []

        tasks = [evaluate_test_case(
            test_case, metric.__class__, metric.threshold,
            metric.model if hasattr(metric, 'model') else "gpt-4o"
        ) for test_case in test_cases]

        metric_results = await asyncio.gather(*tasks, return_exceptions=True)

        valid_scores = [float(result["score"]) for result in metric_results if "score" in result and result["score"] is not None]
        print(f"Valid scores: {valid_scores}")
        avg_score = sum(valid_scores) / len(valid_scores) if valid_scores else 0
        pass_rate = sum(1 for result in metric_results if result.get("passed", False)) / len(metric_results) if metric_results else 0

        results[metric_name] = {
            "average_score": avg_score,
            "pass_rate": pass_rate,
            "individual_results": metric_results
        }
        print(f"{metric_name} - Average Score: {avg_score:.4f}, Pass Rate: {pass_rate:.2%}")

    return results

# if __name__ == "__main__":
def run_rag_qa_golden(file_path, retriever):
    # Generate RAG responses
    rag_response, retrieved_contexts, qa_data = generate_rag_responses_goldenqa(file_path,retriever)
    # Evaluate responses (run repeatedly for debugging)
    eval_results_avg = asyncio.run(evaluate_saved_responses(rag_response, retrieved_contexts, qa_data, file_path))
    # Return evaluation results
    return eval_results_avg
